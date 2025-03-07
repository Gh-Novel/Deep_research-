import os
import sys
import re
import markdown
import time
import hashlib
from datetime import datetime
from io import StringIO
from html.parser import HTMLParser
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Image, Spacer, ListFlowable, ListItem
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT

# Add parent directory to path to allow imports
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from config import EXPORTS_DIR

class MLStripper(HTMLParser):
    """HTML tag stripper for cleaning markdown-to-HTML conversions"""
    def __init__(self):
        super().__init__()
        self.reset()
        self.strict = False
        self.convert_charrefs = True
        self.text = StringIO()
    
    def handle_data(self, d):
        self.text.write(d)
    
    def get_data(self):
        return self.text.getvalue()

def strip_tags(html):
    """Remove HTML tags from text"""
    s = MLStripper()
    s.feed(html)
    return s.get_data()

class ExportAgent:
    def __init__(self):
        """Initialize the export agent with document styles"""
        # Initialize styles for each instance
        self.styles = getSampleStyleSheet()
        self._create_styles()  # Always create styles on initialization

    def _create_styles(self):
        """Create advanced styles with standard ReportLab fonts"""
        # Title style
        self.styles.add(ParagraphStyle(
            name='ReportTitle',
            fontSize=20,
            fontName='Helvetica-Bold',
            textColor=colors.darkblue,
            alignment=TA_CENTER,
            spaceAfter=20,
            spaceBefore=10
        ))
        
        # Heading styles
        self.styles.add(ParagraphStyle(
            name='ReportHeading1',
            fontSize=16,
            fontName='Helvetica-Bold',
            textColor=colors.black,
            alignment=TA_LEFT,
            spaceAfter=16,
            spaceBefore=8
        ))
        
        self.styles.add(ParagraphStyle(
            name='ReportHeading2',
            fontSize=14,
            fontName='Helvetica-Bold',
            textColor=colors.navy,
            alignment=TA_LEFT,
            spaceAfter=14,
            spaceBefore=7
        ))
        
        self.styles.add(ParagraphStyle(
            name='ReportHeading3',
            fontSize=12,
            fontName='Helvetica-Bold',
            textColor=colors.darkblue,
            alignment=TA_LEFT,
            spaceAfter=12,
            spaceBefore=6
        ))
        
        # Body text style
        self.styles.add(ParagraphStyle(
            name='ReportBody',
            fontSize=11,
            fontName='Helvetica',
            textColor=colors.black,
            alignment=TA_LEFT,
            spaceAfter=8,
            spaceBefore=4,
            leading=14
        ))
        
        # Caption style
        self.styles.add(ParagraphStyle(
            name='ChartCaption',
            fontSize=10,
            fontName='Helvetica',
            textColor=colors.darkgrey,
            alignment=TA_CENTER,
            spaceAfter=10,
            spaceBefore=5
        ))

    def export_pdf(self, content, images=None, filename=None):
        """
        Generate PDF with markdown parsing
        
        Args:
            content (str): Markdown content for the report
            images (list): Paths to images to include
            filename (str): Output filename (optional)
            
        Returns:
            str: Path to the generated PDF file
        """
        try:
            # Create exports directory if it doesn't exist
            os.makedirs(EXPORTS_DIR, exist_ok=True)
            
            # Generate a unique filename if not provided
            if not filename:
                # Extract title from content
                title = self._extract_title_from_content(content)
                # Create a timestamp
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                # Create a short hash of the content to ensure uniqueness
                content_hash = hashlib.md5(content.encode()).hexdigest()[:8]
                # Combine to create a unique filename
                filename = f"{title}_{timestamp}_{content_hash}.pdf"
                # Clean the filename to remove invalid characters
                filename = re.sub(r'[\\/*?:"<>|]', "_", filename)
                # Limit filename length
                if len(filename) > 100:
                    filename = filename[:90] + "_" + filename[-9:]
            
            filepath = os.path.join(EXPORTS_DIR, filename)
            
            # Define page size and margins
            page_width, page_height = letter
            margin = 72  # 1 inch margin in points
            
            doc = SimpleDocTemplate(filepath, pagesize=letter, 
                                rightMargin=margin, leftMargin=margin,
                                topMargin=margin, bottomMargin=margin)
            
            story = []
            
            # Title
            story.append(Paragraph("Stock Analysis Report", self.styles['ReportTitle']))
            story.append(Spacer(1, 0.3*inch))
            
            # Process content
            # First, replace markdown formatting with HTML
            # Replace bold text
            content = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', content)
            # Replace italic text
            content = re.sub(r'\*(.*?)\*', r'<i>\1</i>', content)
            
            # Split content into paragraphs
            paragraphs = content.split('\n\n')
            
            for para in paragraphs:
                para = para.strip()
                if not para:
                    continue
                
                # Handle headings
                if para.startswith('# '):
                    heading_text = para[2:].strip()
                    story.append(Paragraph(heading_text, self.styles['ReportHeading1']))
                elif para.startswith('## '):
                    heading_text = para[3:].strip()
                    story.append(Paragraph(heading_text, self.styles['ReportHeading2']))
                elif para.startswith('### '):
                    heading_text = para[4:].strip()
                    story.append(Paragraph(heading_text, self.styles['ReportHeading3']))
                else:
                    # Handle bullet lists
                    if '\n* ' in para or para.startswith('* '):
                        # Split into bullet points
                        bullet_items = []
                        for line in para.split('\n'):
                            line = line.strip()
                            if line.startswith('* '):
                                bullet_text = line[2:].strip()
                                bullet_items.append(Paragraph(bullet_text, self.styles['ReportBody']))
                        
                        if bullet_items:
                            bullet_list = ListFlowable(
                                bullet_items,
                                bulletType='bullet',
                                leftIndent=20,
                                spaceBefore=10,
                                spaceAfter=10
                            )
                            story.append(bullet_list)
                    else:
                        # Regular paragraph
                        story.append(Paragraph(para, self.styles['ReportBody']))
            
            # Images with captions
            if images:
                story.append(Spacer(1, 0.2*inch))
                story.append(Paragraph("Generated Charts", self.styles['ReportHeading2']))
                
                for i, img_path in enumerate(images):
                    if os.path.exists(img_path):
                        try:
                            # Get image dimensions to calculate appropriate size
                            from PIL import Image as PILImage
                            with PILImage.open(img_path) as pil_img:
                                img_width, img_height = pil_img.size
                                
                            # Calculate scaling factor to fit within page
                            content_width = page_width - (2 * margin)
                            width_ratio = content_width / img_width
                            # Limit to 5 inches height at most to ensure it fits on page
                            max_height = 5 * inch
                            
                            # Calculate actual dimensions to use
                            pdf_img_width = min(content_width, img_width * width_ratio)
                            pdf_img_height = min(max_height, img_height * width_ratio)
                            
                            story.append(Spacer(1, 0.2*inch))
                            img = Image(img_path, width=pdf_img_width, height=pdf_img_height)
                            story.append(img)
                        except Exception as e:
                            print(f"Error processing image {img_path}: {str(e)}")
                            # Use a safe fixed size if any issues
                            story.append(Spacer(1, 0.2*inch))
                            img = Image(img_path, width=5*inch, height=3*inch)
                            story.append(img)
                        
                        # Simple text caption without using HTML
                        caption_text = f"Chart {i+1}: {os.path.basename(img_path).replace('_', ' ').replace('.png', '')}"
                        story.append(Paragraph(caption_text, self.styles['ChartCaption']))
            
            # Add footer with page numbers
            def add_page_number(canvas, doc):
                canvas.saveState()
                canvas.setFont('Helvetica', 9)
                page_num_text = f"Page {doc.page}"
                canvas.drawRightString(page_width - margin, margin/2, page_num_text)
                canvas.restoreState()
            
            # Build the document with page numbers
            doc.build(story, onFirstPage=add_page_number, onLaterPages=add_page_number)
            return filepath
            
        except Exception as e:
            import traceback
            print(f"PDF Export Error: {str(e)}")
            print(traceback.format_exc())
            return f"Error generating PDF: {str(e)}"

    def export_word(self, content, images=None):
        """
        Export report to Word document
        
        Args:
            content (str): Markdown content for the report
            images (list): Paths to images to include
            
        Returns:
            str: Path to the generated Word file
        """
        try:
            from docx import Document
            from docx.shared import Inches
            
            # Create exports directory if it doesn't exist
            os.makedirs(EXPORTS_DIR, exist_ok=True)
            
            # Generate a unique filename
            # Extract title from content
            title = self._extract_title_from_content(content)
            # Create a timestamp
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            # Create a short hash of the content to ensure uniqueness
            content_hash = hashlib.md5(content.encode()).hexdigest()[:8]
            # Combine to create a unique filename
            filename = f"{title}_{timestamp}_{content_hash}.docx"
            # Clean the filename to remove invalid characters
            filename = re.sub(r'[\\/*?:"<>|]', "_", filename)
            # Limit filename length
            if len(filename) > 100:
                filename = filename[:90] + "_" + filename[-9:]
            
            filepath = os.path.join(EXPORTS_DIR, filename)
            
            doc = Document()
            doc.add_heading('Stock Analysis Report', 0)
            
            # Add content paragraphs
            paragraphs = content.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    if para.startswith('#'):
                        # Handle markdown headings
                        level = 1
                        while para.startswith('#'):
                            para = para[1:]
                            level += 1
                        doc.add_heading(para.strip(), level=min(level, 9))
                    else:
                        # Handle bullet points
                        if '\n* ' in para or para.startswith('* '):
                            for line in para.split('\n'):
                                line = line.strip()
                                if line.startswith('* '):
                                    bullet_text = line[2:].strip()
                                    # Replace markdown formatting
                                    bullet_text = re.sub(r'\*\*(.*?)\*\*', r'\1', bullet_text)
                                    bullet_text = re.sub(r'\*(.*?)\*', r'\1', bullet_text)
                                    doc.add_paragraph(bullet_text, style='ListBullet')
                        else:
                            # Regular paragraph
                            # Replace markdown formatting
                            para = re.sub(r'\*\*(.*?)\*\*', r'\1', para)
                            para = re.sub(r'\*(.*?)\*', r'\1', para)
                            doc.add_paragraph(para)
            
            # Add images
            if images:
                doc.add_heading('Generated Charts', level=1)
                for img_path in images:
                    if os.path.exists(img_path):
                        try:
                            # Calculate appropriate image width based on page width
                            # Word document default page width is about 6 inches
                            max_width = Inches(5.5)  # Slightly less than page width to ensure margins
                            
                            # Get image dimensions to calculate appropriate size
                            from PIL import Image as PILImage
                            with PILImage.open(img_path) as pil_img:
                                img_width, img_height = pil_img.size
                                aspect_ratio = img_height / img_width
                                
                            # Use a width that fits the page
                            doc.add_picture(img_path, width=max_width)
                        except Exception:
                            # Fallback if PIL is not available or any other error
                            doc.add_picture(img_path, width=Inches(5))
                            
                        caption = os.path.basename(img_path).replace('_', ' ').replace('.png', '')
                        doc.add_paragraph(caption, style='Caption')
            
            # Save the document
            doc.save(filepath)
            return filepath
            
        except Exception as e:
            import traceback
            print(f"Word Export Error: {str(e)}")
            print(traceback.format_exc())
            return f"Error generating Word document: {str(e)}"

    def _extract_title_from_content(self, content):
        """
        Extract a title from the content for use in filenames
        
        Args:
            content (str): The markdown content
            
        Returns:
            str: A sanitized title string
        """
        # Try to find a heading
        heading_match = re.search(r'# (.*?)(\n|$)', content)
        if heading_match:
            title = heading_match.group(1).strip()
        else:
            # If no heading, use the first line or a default
            first_line = content.split('\n', 1)[0].strip()
            title = first_line if first_line else "Report"
        
        # Sanitize the title for use in a filename
        title = re.sub(r'[^a-zA-Z0-9_\-]', '_', title)
        title = re.sub(r'_+', '_', title)  # Replace multiple underscores with a single one
        title = title[:50]  # Limit length
        
        return title 